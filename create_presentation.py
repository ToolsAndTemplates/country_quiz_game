#!/usr/bin/env python3
"""
Script to create a Word presentation document for the Country Quiz Game project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a styled paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para

def create_presentation():
    """Create the Word presentation document."""
    doc = Document()

    # ============================================
    # TITLE PAGE
    # ============================================

    # Add some space at the top
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_heading('Country Quiz Game', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('An Interactive Geography Learning Experience')
    subtitle_run.font.size = Pt(18)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Project type
    type_para = doc.add_paragraph()
    type_run = type_para.add_run('Web Application Project')
    type_run.font.size = Pt(14)
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Technology badge
    tech_para = doc.add_paragraph()
    tech_run = tech_para.add_run('Built with Next.js 14 | React 18 | TypeScript | Tailwind CSS')
    tech_run.font.size = Pt(12)
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break
    doc.add_page_break()

    # ============================================
    # TABLE OF CONTENTS
    # ============================================

    add_heading_with_style(doc, 'Table of Contents', 1)

    toc_items = [
        '1. Project Overview',
        '2. Key Features',
        '3. Technology Stack',
        '4. Game Modes',
        '5. Gamification System',
        '6. User Interface & Design',
        '7. Technical Architecture',
        '8. File Structure',
        '9. Progressive Web App Features',
        '10. Summary'
    ]

    for item in toc_items:
        para = doc.add_paragraph(item)
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================================
    # 1. PROJECT OVERVIEW
    # ============================================

    add_heading_with_style(doc, '1. Project Overview', 1)

    add_heading_with_style(doc, 'What is Country Quiz Game?', 2)

    doc.add_paragraph(
        'The Country Quiz Game is a modern, interactive web application designed to test '
        'and improve geography knowledge through engaging quiz experiences. It combines '
        'educational content with gamification elements to make learning about countries, '
        'capitals, flags, and populations fun and accessible.'
    )

    add_heading_with_style(doc, 'Project Goals', 2)

    goals = [
        'Make geography learning engaging and interactive',
        'Provide multiple quiz modes to test different knowledge areas',
        'Implement gamification to encourage continued learning',
        'Ensure accessibility across all devices (desktop, tablet, mobile)',
        'Enable offline play through Progressive Web App technology'
    ]

    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    add_heading_with_style(doc, 'Target Audience', 2)

    doc.add_paragraph(
        'Students, geography enthusiasts, trivia lovers, and anyone looking to expand '
        'their knowledge of world geography in an entertaining way.'
    )

    doc.add_page_break()

    # ============================================
    # 2. KEY FEATURES
    # ============================================

    add_heading_with_style(doc, '2. Key Features', 1)

    features = [
        ('Three Quiz Modes', 'Flag identification, capital cities matching, and population comparison quizzes'),
        ('Achievement System', '8 unlockable achievements to reward player progress'),
        ('Statistics Dashboard', 'Track your performance with detailed statistics'),
        ('Streak System', 'Build consecutive correct answer streaks for bonus excitement'),
        ('Offline Support', 'Play without internet connection using PWA technology'),
        ('Responsive Design', 'Beautiful experience on any device size'),
        ('Animated UI', 'Smooth animations and micro-interactions throughout'),
        ('Haptic Feedback', 'Physical feedback on mobile devices for correct/incorrect answers'),
        ('195+ Countries', 'Comprehensive country database from REST Countries API'),
        ('Smart Questions', 'Intelligent question generation with regional diversity')
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Feature'
    header_cells[1].text = 'Description'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for feature, description in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = description

    doc.add_page_break()

    # ============================================
    # 3. TECHNOLOGY STACK
    # ============================================

    add_heading_with_style(doc, '3. Technology Stack', 1)

    add_heading_with_style(doc, 'Frontend Framework', 2)

    frontend_tech = [
        ('Next.js 14', 'React framework with App Router for routing and server-side rendering'),
        ('React 18.3', 'UI library for component-based development'),
        ('TypeScript 5', 'Type-safe JavaScript for better code quality')
    ]

    for tech, desc in frontend_tech:
        para = doc.add_paragraph()
        run = para.add_run(f'{tech}: ')
        run.bold = True
        para.add_run(desc)

    add_heading_with_style(doc, 'Styling & Animation', 2)

    styling_tech = [
        ('Tailwind CSS 3.4', 'Utility-first CSS framework for rapid styling'),
        ('Framer Motion 11', 'Advanced animations and micro-interactions'),
        ('PostCSS & Autoprefixer', 'CSS preprocessing and browser compatibility')
    ]

    for tech, desc in styling_tech:
        para = doc.add_paragraph()
        run = para.add_run(f'{tech}: ')
        run.bold = True
        para.add_run(desc)

    add_heading_with_style(doc, 'Data & APIs', 2)

    doc.add_paragraph(
        'REST Countries API (v3.1): Provides real-time data for 195+ countries including '
        'names, capitals, populations, flags, and regional information. The application '
        'implements server-side caching with 1-hour expiration to optimize API usage.'
    )

    add_heading_with_style(doc, 'Development Tools', 2)

    dev_tools = ['ESLint for code quality', 'TypeScript strict mode', 'Path aliasing for clean imports']
    for tool in dev_tools:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # 4. GAME MODES
    # ============================================

    add_heading_with_style(doc, '4. Game Modes', 1)

    add_heading_with_style(doc, 'Flag Master Mode', 2)
    doc.add_paragraph(
        'Players are shown a country flag and must identify the correct country name '
        'from four multiple-choice options. The game uses smart selection to favor '
        'well-known countries and generates wrong answers from similar regions to '
        'increase challenge.'
    )

    add_heading_with_style(doc, 'Capital Cities Mode', 2)
    doc.add_paragraph(
        'A country name is displayed, and players must select the correct capital city '
        'from four options. Only countries with complete capital data are included. '
        'This mode tests knowledge of world capitals and their associated countries.'
    )

    add_heading_with_style(doc, 'Population Quiz Mode', 2)
    doc.add_paragraph(
        'Two countries are shown side-by-side, and players must determine which has '
        'the larger population. This binary choice format provides a different style '
        'of gameplay while teaching about relative country sizes.'
    )

    add_heading_with_style(doc, 'Game Flow', 2)

    flow_steps = [
        'Select quiz mode from home screen',
        'Answer 10 questions per quiz',
        'Receive immediate feedback (1.5 seconds) after each answer',
        'Experience haptic vibration on mobile devices',
        'View results screen with score and personalized message',
        'Stats automatically saved and achievements checked'
    ]

    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ============================================
    # 5. GAMIFICATION SYSTEM
    # ============================================

    add_heading_with_style(doc, '5. Gamification System', 1)

    add_heading_with_style(doc, 'Achievement System (8 Total)', 2)

    achievements = [
        ('First Steps', 'Complete your first quiz'),
        ('Perfect Score', 'Get 10/10 in a quiz'),
        ('Veteran', 'Complete 10 quizzes'),
        ('Expert', 'Complete 50 quizzes'),
        ('On Fire', 'Get 5 consecutive correct answers'),
        ('Unstoppable', 'Get 10 consecutive correct answers'),
        ('Scholar', 'Score 8 or more in a quiz'),
        ('Genius', 'Score 9 or more in a quiz')
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Achievement'
    header_cells[1].text = 'Requirement'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True

    for achievement, requirement in achievements:
        row_cells = table.add_row().cells
        row_cells[0].text = achievement
        row_cells[1].text = requirement

    doc.add_paragraph()

    add_heading_with_style(doc, 'Streak System', 2)
    doc.add_paragraph(
        'The streak system tracks consecutive correct answers. When a streak of 2 or more '
        'is achieved, a visual indicator appears that scales and pulses with flame animations. '
        'The best streak is saved across sessions, encouraging players to beat their record.'
    )

    add_heading_with_style(doc, 'Statistics Dashboard', 2)

    stats = [
        'Total games played across all modes',
        'Cumulative score from all quizzes',
        'High scores per quiz mode (flags, capitals, population)',
        'Best streak (all-time consecutive correct answers)',
        'Achievement progress (X/8 unlocked)'
    ]

    for stat in stats:
        doc.add_paragraph(stat, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # 6. USER INTERFACE & DESIGN
    # ============================================

    add_heading_with_style(doc, '6. User Interface & Design', 1)

    add_heading_with_style(doc, 'Design Philosophy', 2)
    doc.add_paragraph(
        'The application employs a modern glassmorphism design language with animated '
        'gradient backgrounds, creating a visually appealing and immersive experience. '
        'The design is mobile-first, ensuring optimal usability on all screen sizes.'
    )

    add_heading_with_style(doc, 'Visual Elements', 2)

    visuals = [
        ('Glassmorphism', 'Semi-transparent cards with blur effects and subtle borders'),
        ('Animated Gradient', '4-color gradient background with smooth animation'),
        ('Framer Motion', 'Page transitions, button hover effects, streak pulses'),
        ('Floating Particles', 'Background particles for added visual depth'),
        ('Progress Bar', 'Animated question progress indicator')
    ]

    for element, desc in visuals:
        para = doc.add_paragraph()
        run = para.add_run(f'{element}: ')
        run.bold = True
        para.add_run(desc)

    add_heading_with_style(doc, 'Micro-Interactions', 2)

    interactions = [
        'Button scale on hover (1.05x) and press (0.95x)',
        'Question slide animations (left/right entrance)',
        'Achievement pop-in with cascade delay',
        'Icon rotation effects',
        'Loading spinners with animation'
    ]

    for interaction in interactions:
        doc.add_paragraph(interaction, style='List Bullet')

    add_heading_with_style(doc, 'Haptic Feedback', 2)
    doc.add_paragraph(
        'On devices that support vibration, the app provides tactile feedback: '
        'a short 50ms pulse for correct answers, and a pattern (100ms-50ms-100ms) '
        'for incorrect answers, enhancing the mobile gaming experience.'
    )

    doc.add_page_break()

    # ============================================
    # 7. TECHNICAL ARCHITECTURE
    # ============================================

    add_heading_with_style(doc, '7. Technical Architecture', 1)

    add_heading_with_style(doc, 'Component Architecture', 2)
    doc.add_paragraph(
        'The application follows a component-based architecture with clear separation '
        'of concerns. Reusable components are stored in the /components directory, '
        'utility functions in /lib, and type definitions in /types.'
    )

    add_heading_with_style(doc, 'Key Components', 2)

    components = [
        ('QuizContainer', 'Core game engine handling question progression, scoring, and feedback'),
        ('StatsPanel', 'Statistics dashboard modal displaying player progress'),
        ('PWAInstall', 'Progressive Web App installation prompt component')
    ]

    for comp, desc in components:
        para = doc.add_paragraph()
        run = para.add_run(f'{comp}: ')
        run.bold = True
        para.add_run(desc)

    add_heading_with_style(doc, 'State Management', 2)
    doc.add_paragraph(
        'The application uses React hooks (useState, useEffect) for component state '
        'and localStorage for persistent data storage. Game statistics, achievements, '
        'and high scores are automatically saved to the browser\'s local storage.'
    )

    add_heading_with_style(doc, 'API Integration', 2)
    doc.add_paragraph(
        'A Next.js API route (/api/countries) handles communication with the REST '
        'Countries API. Server-side caching with 1-hour expiration optimizes performance '
        'and reduces external API calls. The client also caches questions in memory.'
    )

    add_heading_with_style(doc, 'Error Handling', 2)

    error_features = [
        'Graceful error screens with retry functionality',
        'Fallback to cached data when API fails',
        'User-friendly error messages',
        'Loading states with animated indicators'
    ]

    for feature in error_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # 8. FILE STRUCTURE
    # ============================================

    add_heading_with_style(doc, '8. File Structure', 1)

    structure = """
/country_quiz_game
├── /app                        # Next.js App Router
│   ├── page.tsx               # Home page with quiz mode selection
│   ├── layout.tsx             # Root layout with PWA setup
│   ├── globals.css            # Global styles
│   └── /quiz                  # Quiz mode pages
│       ├── /flags/page.tsx    # Flag identification quiz
│       ├── /capitals/page.tsx # Capital city quiz
│       └── /population/page.tsx # Population quiz
│   └── /api/countries/route.ts # Backend API route
│
├── /components                 # Reusable React components
│   ├── QuizContainer.tsx      # Main quiz logic
│   ├── StatsPanel.tsx         # Statistics dashboard
│   └── PWAInstall.tsx         # PWA installation prompt
│
├── /lib                        # Utility functions
│   ├── countries.ts           # Question generation
│   └── storage.ts             # LocalStorage management
│
├── /types                      # TypeScript interfaces
│   └── country.ts             # Type definitions
│
├── /public                     # Static assets
│   ├── manifest.json          # PWA manifest
│   ├── sw.js                  # Service Worker
│   └── /icons/                # App icons
│
└── Configuration Files
    ├── package.json           # Dependencies
    ├── tsconfig.json          # TypeScript config
    ├── tailwind.config.ts     # Tailwind config
    └── next.config.js         # Next.js config
"""

    para = doc.add_paragraph()
    run = para.add_run(structure)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ============================================
    # 9. PROGRESSIVE WEB APP FEATURES
    # ============================================

    add_heading_with_style(doc, '9. Progressive Web App Features', 1)

    add_heading_with_style(doc, 'Offline Capability', 2)
    doc.add_paragraph(
        'The application includes a Service Worker that enables offline play. '
        'Once installed, users can enjoy the quiz even without an internet connection, '
        'making it perfect for learning on-the-go.'
    )

    add_heading_with_style(doc, 'Installation', 2)
    doc.add_paragraph(
        'Users can install the app to their device\'s home screen for a native-like '
        'experience. The PWAInstall component detects installation eligibility and '
        'prompts users to add the app to their device.'
    )

    add_heading_with_style(doc, 'App Shortcuts', 2)
    doc.add_paragraph(
        'The PWA manifest includes shortcuts that allow users to launch directly '
        'into specific quiz modes from their device\'s home screen or app launcher.'
    )

    add_heading_with_style(doc, 'PWA Features Summary', 2)

    pwa_features = [
        'Service Worker for asset caching and offline support',
        'Web App Manifest with app metadata and icons',
        'Standalone display mode (no browser chrome)',
        'Custom theme colors matching app design',
        'Maskable icons for Android adaptive icons',
        'Quick-launch shortcuts to quiz modes'
    ]

    for feature in pwa_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # 10. SUMMARY
    # ============================================

    add_heading_with_style(doc, '10. Summary', 1)

    doc.add_paragraph(
        'The Country Quiz Game is a comprehensive, modern web application that '
        'demonstrates best practices in React development with Next.js. It combines '
        'educational content with engaging gamification to create an enjoyable '
        'geography learning experience.'
    )

    add_heading_with_style(doc, 'Technical Highlights', 2)

    highlights = [
        'Built with Next.js 14 App Router and React 18',
        'Full TypeScript implementation with strict mode',
        'Tailwind CSS for responsive, utility-first styling',
        'Framer Motion for smooth animations',
        'Progressive Web App with offline support',
        'REST Countries API integration with smart caching',
        'Comprehensive gamification system'
    ]

    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')

    add_heading_with_style(doc, 'Key Statistics', 2)

    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'

    header_cells = stats_table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True

    stats_data = [
        ('Quiz Modes', '3'),
        ('Countries in Database', '195+'),
        ('Achievements', '8'),
        ('Questions per Quiz', '10'),
        ('Technology Stack', 'Next.js, React, TypeScript, Tailwind')
    ]

    for metric, value in stats_data:
        row_cells = stats_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    doc.add_paragraph()

    # Final note
    final = doc.add_paragraph()
    final_run = final.add_run(
        'This project showcases the power of modern web technologies to create '
        'engaging, educational experiences that work seamlessly across all devices.'
    )
    final_run.italic = True

    # Save the document
    doc.save('/home/user/country_quiz_game/Country_Quiz_Game_Presentation.docx')
    print('Presentation created successfully: Country_Quiz_Game_Presentation.docx')

if __name__ == '__main__':
    create_presentation()
